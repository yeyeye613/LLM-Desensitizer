
import json
import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import letter

# Create output directory
output_dir = "test_data"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# 1. Generate DOCX (Resume Scenario)
def create_docx():
    doc = Document()
    doc.add_heading('Resume - Li Ming', 0)

    doc.add_paragraph('Prompt to LLM: "Please review my resume and suggest improvements for a senior developer role."')
    
    doc.add_heading('Contact Information', level=1)
    doc.add_paragraph('Name: Li Ming')
    doc.add_paragraph('Phone: 13912345678')
    doc.add_paragraph('Email: liming_dev@example.com')
    doc.add_paragraph('Address: Room 402, Unit 1, Building 5, West Lake District, Hangzhou')
    
    doc.add_heading('Personal Details', level=1)
    doc.add_paragraph('ID Card: 330106198501011234')
    doc.add_paragraph('Date of Birth: 1985-01-01')
    
    doc.add_heading('Experience', level=1)
    p = doc.add_paragraph('Senior Java Developer at Tech Corp (2018-Present). Led the development of the payment gateway using Spring Boot.')
    
    filename = os.path.join(output_dir, "user_resume_prompt.docx")
    doc.save(filename)
    print(f"Created {filename}")

# 2. Generate PDF (Debug Scenario)
def create_pdf():
    filename = os.path.join(output_dir, "user_debug_prompt.pdf")
    c = canvas.Canvas(filename, pagesize=letter)
    
    # Simple text drawing
    c.setFont("Helvetica", 12)
    y = 750
    
    lines = [
        "User Prompt: 'My application is failing to connect to the database. Here are the logs and config. Can you help?'",
        "",
        "--- application.properties ---",
        "spring.application.name=PaymentService",
        "server.port=8080",
        "",
        "# Database Configuration",
        "spring.datasource.url=jdbc:mysql://192.168.1.200:3306/payment_db",
        "spring.datasource.username=admin_user",
        "spring.datasource.password=P@ssw0rd123!  <-- Is this correct?",
        "",
        "# Third-party Integrations",
        "api.google.maps.key=AIzaSyD-1234567890abcdef1234567890abcdef",
        "api.stripe.secret=sk_live_51HvT23...",
        "",
        "--- Error Logs ---",
        "2023-10-27 10:00:01 ERROR Connection failed to 192.168.1.200",
        "User 'admin_user' access denied."
    ]
    
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
        
    c.save()
    print(f"Created {filename}")

# 3. Generate JSON (Data Analysis Scenario)
def create_json():
    data = {
        "prompt": "Analyze the following transaction records and identify any potential fraud patterns. Also summarize the top spenders.",
        "dataset": [
            {
                "transaction_id": "TXN10001",
                "user_id": "U12345",
                "name": "Wang Wei",
                "credit_card": "4532711234567895",
                "amount": 5000.00,
                "merchant": "Apple Store",
                "ip_address": "202.106.0.20"
            },
            {
                "transaction_id": "TXN10002",
                "user_id": "U67890",
                "name": "Liu Fang",
                "credit_card": "6222021234567890",
                "amount": 120.50,
                "merchant": "Starbucks",
                "ip_address": "192.168.1.5"
            },
            {
                "transaction_id": "TXN10003",
                "user_id": "U11223",
                "name": "Zhang San",
                "phone": "13800138000",
                "note": "Customer requested refund to bank account 6217001234567890",
                "amount": -50.00
            }
        ],
        "system_instruction": "Do not output full credit card numbers in the summary."
    }
    
    filename = os.path.join(output_dir, "user_analysis_prompt.json")
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Created {filename}")

if __name__ == "__main__":
    try:
        create_docx()
        create_pdf()
        create_json()
        print("All files generated successfully.")
    except Exception as e:
        print(f"Error generating files: {e}")
